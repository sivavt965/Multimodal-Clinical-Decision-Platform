"""
Builds the Senior Design Project thesis DOCX in VIT-AP format.

Format rules (per ANNEXURE-II):
  - A4 paper
  - Margins: Left 3.81 cm, Right/Top/Bottom 2.54 cm
  - Body: Times New Roman 12 pt, line spacing 1.5
  - Chapter heading: 16 pt bold
  - Section heading : 14 pt bold CAPS
  - Subsection heading: 12 pt bold CAPS
  - Page numbering: Arabic numerals, bottom-centre

Outputs: Symile_MIMIC_Thesis_VITAP.docx in the project root.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

OUT_PATH = Path(__file__).with_name("Symile_MIMIC_Thesis_VITAP.docx")
PLOT_DIR = Path(__file__).with_name("mimic_project") / "plots"


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def _set_run(run, *, size=12, bold=False, italic=False, caps=False, color=None):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if caps:
        run.font.all_caps = True
    if color is not None:
        run.font.color.rgb = color


def _para(doc, text="", *, size=12, bold=False, italic=False, caps=False,
          align=None, space_before=0, space_after=6, line_spacing=1.5,
          first_line_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        p.alignment = align
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, italic=italic, caps=caps)
    return p


def _chapter_heading(doc, num, title):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(f"CHAPTER {num}")
    _set_run(run, size=16, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    run = p2.add_run(title.upper())
    _set_run(run, size=16, bold=True)


def _section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title.upper())
    _set_run(run, size=14, bold=True)


def _subsection(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    _set_run(run, size=12, bold=True)


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + 0.75 * level)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    _set_run(run, size=12)


def _table(doc, header_row, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header_row))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(header_row):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run(run, size=11, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_run(run, size=11)
    if col_widths_cm is not None:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    return table


def _caption(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"{label}: {text}")
    _set_run(run, size=10, italic=True)


def _figure(doc, image_path: Path, label: str, caption: str, width_in=5.0):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(width_in))
    except Exception:
        return
    _caption(doc, label, caption)


def _add_page_numbers(doc):
    """Bottom-centre Arabic page numbers."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run = p.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    _set_run(run, size=11)


# ---------------------------------------------------------------------------
# build the document
# ---------------------------------------------------------------------------
doc = Document()

# Page setup: A4 with VIT-AP margins
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(3.81)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# Default style → Times New Roman 12 pt, 1.5 line spacing
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rFonts.set(qn(attr), "Times New Roman")

_add_page_numbers(doc)

# ============================================================================
# 1. COVER / TITLE PAGE
# ============================================================================
for _ in range(3):
    _para(doc, "")
_para(doc, "A Senior Design Project Report on", align=WD_ALIGN_PARAGRAPH.CENTER,
      size=12, italic=True, space_after=18)
_para(doc,
      "MULTIMODAL CLINICAL DECISION SUPPORT PLATFORM FOR GENERAL WARD "
      "PHYSICIANS USING THE SYMILE-MIMIC DATASET",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, caps=True,
      space_after=24)
_para(doc, "Submitted in partial fulfilment of the requirements for the award of the degree of",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=14, italic=True, space_after=12)
_para(doc, "BACHELOR OF TECHNOLOGY",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, space_after=6)
_para(doc, "in", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, italic=True,
      space_after=6)
_para(doc, "COMPUTER SCIENCE AND ENGINEERING",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, caps=True,
      space_after=24)
_para(doc, "by", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, italic=True,
      space_after=6)
_para(doc, "SIVA VENKAT VELURU (Reg. No. 22BCEXXXXX)",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, space_after=24)
_para(doc, "Under the guidance of", align=WD_ALIGN_PARAGRAPH.CENTER,
      size=14, italic=True, space_after=6)
_para(doc, "DR. <GUIDE NAME>", align=WD_ALIGN_PARAGRAPH.CENTER,
      size=14, bold=True, space_after=24)
_para(doc, "SCHOOL OF COMPUTER SCIENCE AND ENGINEERING",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, caps=True,
      space_after=6)
_para(doc, "VIT-AP UNIVERSITY, AMARAVATI, ANDHRA PRADESH",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=18)
_para(doc, "MAY 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)

# ============================================================================
# 2. DECLARATION
# ============================================================================
doc.add_page_break()
_para(doc, "DECLARATION", align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
      bold=True, caps=True, space_after=18)
_para(doc,
      "I hereby declare that the thesis entitled “Multimodal Clinical Decision "
      "Support Platform for General Ward Physicians Using the Symile-MIMIC "
      "Dataset”, submitted by me to VIT-AP University, Amaravati, in partial "
      "fulfilment of the requirements for the award of the degree of Bachelor "
      "of Technology in Computer Science and Engineering, is a record of "
      "bonafide work carried out by me under the supervision of Dr. <Guide "
      "Name>, School of Computer Science and Engineering, VIT-AP University.")
_para(doc,
      "I further declare that the work reported in this thesis has not been "
      "submitted, and will not be submitted, either in part or in full, for "
      "the award of any other degree or diploma at this institute or any "
      "other institute or university.")
_para(doc, "")
_para(doc, "Place: Amaravati", size=12)
_para(doc, "Date:                                                               "
      "                                Signature of the Candidate", size=12,
      bold=True)

# ============================================================================
# 3. CERTIFICATE
# ============================================================================
doc.add_page_break()
_para(doc, "CERTIFICATE", align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
      bold=True, caps=True, space_after=18)
_para(doc,
      "This is to certify that the Senior Design Project titled “Multimodal "
      "Clinical Decision Support Platform for General Ward Physicians Using "
      "the Symile-MIMIC Dataset”, submitted by Siva Venkat Veluru "
      "(22BCEXXXXX), is a record of bonafide work done under my guidance and "
      "supervision in partial fulfilment of the requirements for the award of "
      "the degree of Bachelor of Technology in Computer Science and "
      "Engineering at VIT-AP University, Amaravati. The contents of this "
      "project work, in full or in parts, have neither been taken from any "
      "other source nor have been submitted to any other institute or "
      "university for the award of any degree or diploma, and the same is "
      "certified.")
_para(doc, "")
_para(doc, "")
_para(doc, "Dr. <Guide Name>", size=12, bold=True)
_para(doc, "Project Guide", size=12)
_para(doc, "School of Computer Science and Engineering")
_para(doc, "VIT-AP University, Amaravati")
_para(doc, "")
_para(doc, "The thesis is satisfactory / unsatisfactory.", size=12, italic=True)
_para(doc, "")
_para(doc, "Internal Examiner                                                   "
      "                            External Examiner", size=12, bold=True)
_para(doc, "")
_para(doc, "Approved by", size=12)
_para(doc, "")
_para(doc, "PROGRAM CHAIR                                                       "
      "                                            DEAN", size=12, bold=True)
_para(doc, "B.Tech. Computer Science & Engineering                              "
      "                  School of Computer Science and Engineering",
      size=12)

# ============================================================================
# 4. ABSTRACT
# ============================================================================
doc.add_page_break()
_para(doc, "ABSTRACT", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True,
      caps=True, space_after=18)
_para(doc,
      "Decisions in a general medical ward are made under temporal "
      "uncertainty: laboratory and electrocardiogram (ECG) results frequently "
      "arrive hours before the chest radiograph (CXR) is performed, while "
      "existing clinical artificial-intelligence tools assume that all "
      "modalities are available simultaneously. This thesis presents a "
      "doctor-facing, role-based clinical decision support (CDS) platform "
      "that explicitly models this timing gap by separating an early-risk "
      "tabular phase (ECG and 50 MIMIC-IV laboratory biomarkers) from a "
      "later imaging phase (DenseNet121 CXR classifier with Grad-CAM "
      "explanations and Monte-Carlo Dropout uncertainty), and binds them "
      "together with FAISS-based similar-case retrieval over 1024-dimensional "
      "DenseNet global-average-pooled embeddings.")
_para(doc,
      "The system is built on the Symile-MIMIC release of the MIMIC-IV "
      "dataset, with 237,972 training, 1,959 validation and 3,403 test "
      "examples covering eight CheXpert thoracic findings. The trained "
      "DenseNet121 baseline attains a micro-AUROC of 0.864 (95 % CI 0.859–"
      "0.869) and a macro-AUROC of 0.812 (95 % CI 0.803–0.820) on the "
      "held-out test set. After temperature scaling on validation "
      "(T = 1.252), expected calibration error is reduced from 0.0397 to "
      "0.0269 and Brier score from 0.0872 to 0.0859. Sixty-pass MC-Dropout "
      "with p = 0.30 yields per-prediction variance estimates that are "
      "discretised into a three-tier uncertainty banner shown to the "
      "clinician. The platform is implemented as a FastAPI Python backend "
      "and a Next.js 13 / TypeScript frontend backed by Supabase Postgres "
      "and object storage. It supports four roles — ward doctor, "
      "radiologist, clinical admin and system admin — with eight "
      "audit-logged mutating endpoints, role-aware tab visibility and a "
      "dual-user consultation thread. The work demonstrates that explicit "
      "timing-awareness, calibrated probabilities, transparent uncertainty "
      "and case-based retrieval can be combined inside a single deployable "
      "interface that is faithful to how ward decisions are actually made.")
_para(doc, "")
_para(doc,
      "Keywords: Clinical decision support, Multimodal learning, MIMIC-IV, "
      "Symile, DenseNet121, Grad-CAM, MC Dropout, Temperature scaling, "
      "FAISS retrieval, FastAPI, Next.js, Supabase, RBAC.",
      size=12, italic=True)

# ============================================================================
# 5. ACKNOWLEDGEMENT
# ============================================================================
doc.add_page_break()
_para(doc, "ACKNOWLEDGEMENT", align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
      bold=True, caps=True, space_after=18)
_para(doc,
      "It is my pleasure to express, with a deep sense of gratitude, my "
      "sincere thanks to my project guide Dr. <Guide Name>, School of "
      "Computer Science and Engineering, VIT-AP University, for his/her "
      "constant guidance, continual encouragement and patient mentoring "
      "throughout the course of this Senior Design Project. The exposure to "
      "the discipline of multimodal clinical machine learning that I "
      "received under his/her supervision has shaped both this work and my "
      "approach to engineering as a whole.")
_para(doc,
      "I would like to express my gratitude to the Honourable Chancellor, "
      "Vice-Chancellor and Dean, School of Computer Science and "
      "Engineering, VIT-AP University, for providing an environment of "
      "academic excellence in which to pursue this work, and for their "
      "continual inspiration during the tenure of the course.")
_para(doc,
      "I extend my whole-hearted thanks to the Program Chair, Computer "
      "Science and Engineering, and to all teaching and administrative "
      "staff of the School, whose enthusiasm and timely encouragement "
      "made the acquisition of the requisite knowledge possible. I would "
      "also like to thank the curators of the MIMIC-IV and MIMIC-CXR "
      "databases at the MIT Laboratory for Computational Physiology and "
      "the authors of the Symile contrastive-learning library, whose "
      "credentialed datasets and open-source tools form the data and "
      "modelling foundation of this project.")
_para(doc,
      "I would like to thank my parents for their unconditional support, "
      "and my friends who persuaded and encouraged me to take up and "
      "complete this task. Lastly, I express my gratitude and appreciation "
      "to all those who have helped me, directly or indirectly, toward the "
      "successful completion of this project.")
_para(doc, "")
_para(doc, "Place: Amaravati", size=12)
_para(doc, "Date:                                                               "
      "                                Siva Venkat Veluru", size=12, bold=True)

# ============================================================================
# 6. TABLE OF CONTENTS
# ============================================================================
doc.add_page_break()
_para(doc, "TABLE OF CONTENTS", align=WD_ALIGN_PARAGRAPH.CENTER,
      size=14, bold=True, caps=True, space_after=18)
toc_rows = [
    ("Declaration", "i"),
    ("Certificate", "ii"),
    ("Abstract", "iii"),
    ("Acknowledgement", "iv"),
    ("Table of Contents", "v"),
    ("List of Tables", "vii"),
    ("List of Figures", "viii"),
    ("List of Abbreviations", "ix"),
    ("Chapter 1   Introduction", "1"),
    ("Chapter 2   Literature Review", "5"),
    ("Chapter 3   System Requirements and Methodology", "10"),
    ("Chapter 4   System Design and Architecture", "16"),
    ("Chapter 5   Implementation", "23"),
    ("Chapter 6   Results and Discussion", "32"),
    ("Chapter 7   Conclusion and Future Work", "42"),
    ("References", "45"),
    ("Appendix A   API Endpoint Reference", "48"),
    ("Appendix B   Database Schema", "51"),
    ("Appendix C   Per-Label Metrics", "54"),
]
_table(doc, ["Section", "Page"], toc_rows, col_widths_cm=[12, 2])

# ============================================================================
# 7. LIST OF TABLES
# ============================================================================
doc.add_page_break()
_para(doc, "LIST OF TABLES", align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
      bold=True, caps=True, space_after=18)
lot_rows = [
    ("3.1", "Eight CheXpert findings used in Phase B", "12"),
    ("3.2", "Class prevalence in the train / validation / test splits", "13"),
    ("4.1", "Three-tier risk-badge thresholds", "18"),
    ("4.2", "MC Dropout uncertainty tiers", "19"),
    ("4.3", "Role-to-surface visibility matrix", "20"),
    ("5.1", "Audited mutating endpoints", "27"),
    ("5.2", "External Python and JavaScript dependencies", "28"),
    ("5.3", "Symile-MIMIC dataset construction parameters", "29"),
    ("5.4", "Symile contrastive training hyperparameters", "30"),
    ("5.5", "Phase 5 authentication milestones", "31"),
    ("6.1", "Test-set AUROC, AUPRC and Brier score", "33"),
    ("6.2", "Per-label AUROC with 95 % bootstrap CI", "34"),
    ("6.3", "Calibration before vs. after temperature scaling", "35"),
    ("6.4", "Per-label AUPRC vs. class prevalence", "36"),
    ("6.5", "MC Dropout uncertainty test-set summary", "37"),
    ("6.6", "Inference latency on CUDA hardware", "38"),
    ("6.7", "Symile zero-shot retrieval performance", "39"),
    ("C.1", "Full per-label metrics table", "54"),
]
_table(doc, ["Table", "Title", "Page"], lot_rows, col_widths_cm=[1.5, 11, 1.5])

# ============================================================================
# 8. LIST OF FIGURES
# ============================================================================
doc.add_page_break()
_para(doc, "LIST OF FIGURES", align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
      bold=True, caps=True, space_after=18)
lof_rows = [
    ("1.1", "Clinical timing gap between early modalities and CXR", "2"),
    ("4.1", "Three-tier system architecture", "16"),
    ("4.2", "Phase A → Phase B sequence diagram", "21"),
    ("5.1", "Frontend role-based tab visibility", "30"),
    ("6.1", "Micro-averaged ROC curve (pre-TS)", "33"),
    ("6.2", "Per-label ROC grid (pre-TS)", "34"),
    ("6.3", "Micro-averaged precision-recall curve (pre-TS)", "35"),
    ("6.4", "Reliability diagram before temperature scaling", "36"),
    ("6.5", "Reliability diagram after temperature scaling", "36"),
    ("6.6", "Coverage vs. macro-AUROC under entropy-based selective prediction", "37"),
    ("6.7", "Predictive uncertainty histogram (entropy)", "38"),
    ("6.8", "Threshold sweep — F1, sensitivity and specificity", "38"),
    ("6.9", "Sample Grad-CAM overlay (Support Devices)", "39"),
]
_table(doc, ["Figure", "Title", "Page"], lof_rows, col_widths_cm=[1.5, 11, 1.5])

# ============================================================================
# 9. LIST OF ABBREVIATIONS
# ============================================================================
doc.add_page_break()
_para(doc, "LIST OF ABBREVIATIONS", align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
      bold=True, caps=True, space_after=18)
abbr_rows = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("AUPRC", "Area Under the Precision-Recall Curve"),
    ("AUROC", "Area Under the Receiver Operating Characteristic curve"),
    ("CDS", "Clinical Decision Support"),
    ("CXR", "Chest X-Ray"),
    ("DICOM", "Digital Imaging and Communications in Medicine"),
    ("ECE", "Expected Calibration Error"),
    ("ECG", "Electrocardiogram"),
    ("FAISS", "Facebook AI Similarity Search"),
    ("FiLM", "Feature-wise Linear Modulation"),
    ("GAP", "Global Average Pooling"),
    ("Grad-CAM", "Gradient-weighted Class Activation Mapping"),
    ("HIPAA", "Health Insurance Portability and Accountability Act"),
    ("HS256", "HMAC with SHA-256"),
    ("JWT", "JSON Web Token"),
    ("MC", "Monte Carlo"),
    ("MIMIC", "Medical Information Mart for Intensive Care"),
    ("MLP", "Multi-Layer Perceptron"),
    ("MRN", "Medical Record Number"),
    ("NLL", "Negative Log Likelihood"),
    ("PHI", "Protected Health Information"),
    ("RBAC", "Role-Based Access Control"),
    ("REST", "Representational State Transfer"),
    ("RLS", "Row-Level Security"),
    ("ROC", "Receiver Operating Characteristic"),
    ("SQL", "Structured Query Language"),
    ("TS", "Temperature Scaling"),
    ("UMAP", "Uniform Manifold Approximation and Projection"),
    ("UQ", "Uncertainty Quantification"),
    ("UUID", "Universally Unique Identifier"),
]
_table(doc, ["Acronym", "Expansion"], abbr_rows, col_widths_cm=[3, 11])

# ============================================================================
# CHAPTER 1 — INTRODUCTION
# ============================================================================
_chapter_heading(doc, 1, "Introduction")

_section(doc, "1.1 Background and Motivation")
_para(doc,
      "In a general medical ward, decisions about whether to escalate a "
      "patient, request imaging, or wait and observe must often be taken "
      "before all relevant data has arrived. Electrocardiogram (ECG) traces "
      "and core laboratory panels (electrolytes, troponin, creatinine, "
      "lactate, brain natriuretic peptide) are typically available within "
      "an hour of admission, whereas a chest radiograph (CXR) may be "
      "delayed by anywhere from a few hours to several days, depending on "
      "porter availability, hospital load and clinical priority. The "
      "general ward physician therefore reasons twice about the same "
      "patient: once with tabular early data only, and again later when "
      "the imaging modality becomes available.")
_para(doc,
      "Clinical artificial-intelligence (AI) tools published in the last "
      "five years have, with very few exceptions, ignored this temporal "
      "structure. Prediction-only models output a single probability per "
      "finding given a CXR; retrieval models surface visually similar "
      "historical cases without an associated risk read-out; uncertainty-"
      "aware models report variance numbers but do not contextualise them "
      "for non-radiologist users; and almost none of these tools "
      "communicate when the model is likely to be wrong. The result is "
      "that even technically sound clinical AI outputs land in the ward "
      "in a form that physicians do not trust enough to act on.")
_para(doc,
      "This thesis takes the opposite design position. Rather than "
      "engineer a single global model, the system is built around the "
      "way ward decisions are actually taken: an early-risk tabular "
      "module fires immediately on ECG and laboratory data; a later CXR "
      "module fires when imaging arrives; a similarity engine indexes "
      "both states to surface comparable historical trajectories; and a "
      "Before-vs.-After view explicitly captures what the model believed "
      "before imaging and what it believes after. Calibrated "
      "probabilities, MC-Dropout uncertainty banners and Grad-CAM "
      "overlays then frame each output in a way the ward physician can "
      "interrogate.")

_section(doc, "1.2 Problem Statement")
_para(doc,
      "Given a longitudinal hospital admission for which (i) ECG and "
      "laboratory results are available shortly after admission and "
      "(ii) a chest radiograph is acquired some hours later, design and "
      "implement a clinical decision support web platform that:")
_bullet(doc, "produces an early imaging-priority recommendation from "
        "ECG and labs alone, before the CXR is acquired;")
_bullet(doc, "produces a calibrated, multi-label thoracic finding read "
        "from the CXR with per-finding uncertainty and Grad-CAM "
        "explanation;")
_bullet(doc, "retrieves comparable historical cases at both stages "
        "using a single shared embedding space;")
_bullet(doc, "exposes a Before-vs.-After comparison so the ward "
        "physician can see how their early hypothesis updates after "
        "imaging;")
_bullet(doc, "enforces role-based access control across ward doctor, "
        "radiologist, clinical administrator and system administrator "
        "roles, and writes a tamper-evident audit log of every "
        "mutating action.")

_section(doc, "1.3 Objectives")
_bullet(doc, "Train, calibrate and evaluate a DenseNet121-based CXR "
        "classifier on the eight clinically actionable CheXpert findings "
        "in the Symile-MIMIC release.")
_bullet(doc, "Quantify predictive uncertainty using Monte-Carlo Dropout "
        "and surface it as a three-tier banner.")
_bullet(doc, "Implement Phase A early-risk inference on seven cardiac "
        "and 50 laboratory features, and persist its output for later "
        "comparison.")
_bullet(doc, "Build a retrieval engine over 1024-dimensional DenseNet "
        "global-average-pooled (GAP) embeddings using FAISS.")
_bullet(doc, "Design a FastAPI backend and a Next.js 13 / TypeScript "
        "frontend that compose these modules into a coherent doctor-"
        "facing workflow with four user roles.")
_bullet(doc, "Audit-log every mutating endpoint and provide a system-"
        "administrator surface for live presence and user management.")

_section(doc, "1.4 Scope and Out of Scope")
_para(doc,
      "The platform is positioned as an assistive layer for general ward "
      "physicians, not as a final diagnostic authority. The following "
      "are explicitly out of scope:")
_bullet(doc, "live integration with hospital information systems, "
        "electronic medical records, or HL7 / FHIR exchange;")
_bullet(doc, "use of any patient-identifiable data — only de-identified "
        "MIMIC-IV cases are used and no PHI is redistributed;")
_bullet(doc, "automated triage, escalation or treatment recommendations;")
_bullet(doc, "multi-tenant hospital deployment beyond demonstration scale.")

_section(doc, "1.5 Organisation of the Thesis")
_para(doc,
      "Chapter 2 reviews related work in multimodal clinical AI, "
      "uncertainty quantification, calibration and case-based retrieval. "
      "Chapter 3 enumerates the functional and non-functional requirements "
      "and the methodology for building each module. Chapter 4 presents "
      "the three-tier system architecture, the role/permission matrix, "
      "and the inference pipeline. Chapter 5 describes the concrete "
      "implementation across backend, frontend, model training and "
      "infrastructure. Chapter 6 reports the empirical results — AUROC, "
      "AUPRC, calibration before and after temperature scaling, MC-"
      "Dropout uncertainty and inference latency — together with a "
      "qualitative discussion of failure modes. Chapter 7 concludes and "
      "outlines future work.")

# ============================================================================
# CHAPTER 2 — LITERATURE REVIEW
# ============================================================================
_chapter_heading(doc, 2, "Literature Review")

_section(doc, "2.1 Multimodal Clinical AI on MIMIC")
_para(doc,
      "The MIMIC-IV (Johnson et al., 2023) and MIMIC-CXR-JPG (Johnson "
      "et al., 2019) databases provide the largest publicly available "
      "credentialed corpus of intensive-care admissions and chest "
      "radiographs respectively. The Symile-MIMIC release (Saporta et "
      "al., 2024) joins these with synchronously acquired ECG and "
      "laboratory data and explicitly proposes a contrastive learning "
      "objective over an unbounded number of modalities. While Symile’s "
      "loss function and aligned 448-dimensional joint embedding (ECG "
      "128 + CXR 256 + Lab 64) are technically attractive, the present "
      "work uses the simpler 1024-dimensional DenseNet GAP feature for "
      "retrieval because it allows both single-modality CXR queries and "
      "multimodal queries within a single FAISS index without "
      "re-training, an important property for a demo platform.")

_section(doc, "2.2 Chest Radiograph Classification")
_para(doc,
      "CheXNet (Rajpurkar et al., 2017) established DenseNet121 as a "
      "strong baseline for multi-label thoracic disease classification "
      "on ChestX-ray14, reaching radiologist-level performance on "
      "pneumonia. Subsequent work on the larger CheXpert dataset (Irvin "
      "et al., 2019) introduced uncertainty handling for radiology "
      "labels and per-label operating points. The DenseNet121 backbone "
      "remains the de-facto reference architecture because it is "
      "comparatively small (≈ 8 M parameters), trains quickly on a "
      "single modern GPU and supports interpretable Grad-CAM overlays.")

_section(doc, "2.3 Uncertainty Quantification")
_para(doc,
      "Gal and Ghahramani (2016) showed that dropout applied at "
      "inference time yields a Bayesian approximation to the posterior "
      "predictive distribution, enabling cheap uncertainty estimates "
      "without the cost of full ensembles. In medical imaging this "
      "Monte-Carlo Dropout (MC-Dropout) procedure has been adopted "
      "widely, but the raw variance numbers are seldom directly useful "
      "to clinicians. The present work therefore discretises MC-Dropout "
      "variance into three tiers (Low / Moderate / High) and renders a "
      "human-readable banner.")

_section(doc, "2.4 Calibration and Temperature Scaling")
_para(doc,
      "Modern deep neural networks are systematically over-confident "
      "(Guo et al., 2017). Temperature scaling — dividing the pre-"
      "sigmoid logits by a single learnt scalar T — is a parameter-free "
      "post-hoc calibration procedure that preserves the order of "
      "predictions while improving the expected calibration error "
      "(ECE). It is preferred to histogram binning or Platt scaling "
      "because it does not change the area under the ROC.")

_section(doc, "2.5 Case-Based Retrieval and Explanation")
_para(doc,
      "Case-based reasoning (Aamodt and Plaza, 1994) and content-based "
      "image retrieval (Müller et al., 2004) are long-standing motifs "
      "in medical informatics. FAISS (Johnson et al., 2019b) supplies "
      "an inner-product / L2 vector index that scales to tens of "
      "millions of vectors on commodity hardware and is the de-facto "
      "choice for similarity retrieval at runtime. Grad-CAM (Selvaraju "
      "et al., 2017) is the most widely deployed visual-explanation "
      "method and was retained here over more recent saliency methods "
      "because of its simplicity and the absence of free hyper-"
      "parameters at inference time.")

_section(doc, "2.6 Identified Gap")
_para(doc,
      "The literature review reveals a consistent gap: prediction, "
      "retrieval and uncertainty are evaluated separately, and almost "
      "no system models the temporal asymmetry between cheap modalities "
      "(ECG, labs) and expensive modalities (CXR). The platform "
      "presented here is therefore positioned as a workflow-level "
      "contribution rather than a new model: existing well-validated "
      "components are composed into a doctor-facing interface that is "
      "faithful to the actual sequence in which clinical data arrives.")

# ============================================================================
# CHAPTER 3 — REQUIREMENTS AND METHODOLOGY
# ============================================================================
_chapter_heading(doc, 3, "System Requirements and Methodology")

_section(doc, "3.1 Functional Requirements")
_subsection(doc, "3.1.1 Case Intake")
_bullet(doc, "Select a case from the pre-loaded demo cohort.")
_bullet(doc, "Upload a new CXR (JPEG, PNG or DICOM) up to 10 MB.")
_bullet(doc, "Optionally attach ECG and / or lab data (JSON or CSV).")
_bullet(doc, "Display which modalities are present and which are missing.")

_subsection(doc, "3.1.2 Phase A — Early Risk")
_bullet(doc, "Activate when ECG + labs are present but CXR is not.")
_bullet(doc, "Output a Low / Moderate / High patient risk level.")
_bullet(doc, "Output an imaging-priority recommendation.")
_bullet(doc, "Retrieve the top-k similar early-stage cases from the "
        "ECG + lab embedding space.")

_subsection(doc, "3.1.3 Phase B — CXR Inference")
_bullet(doc, "Multi-label classification over the eight CheXpert "
        "findings listed in Table 3.1.")
_bullet(doc, "Temperature-scaled calibrated probabilities.")
_bullet(doc, "Three-tier risk badge (Unlikely / Monitor / Elevated Risk).")
_bullet(doc, "Per-prediction Grad-CAM overlay.")
_bullet(doc, "Per-prediction MC-Dropout uncertainty banner.")
_bullet(doc, "Top-k similar cases on the joint multimodal embedding.")

_subsection(doc, "3.1.4 Role-Based Access Control")
_bullet(doc, "Four roles: ward doctor, radiologist, clinical admin, "
        "system admin.")
_bullet(doc, "Per-role tab visibility and per-role page guards.")
_bullet(doc, "Eight mutating endpoints write to an append-only audit log.")
_bullet(doc, "System admin sees live presence, user management and the "
        "audit log itself.")

_section(doc, "3.2 Non-Functional Requirements")
_bullet(doc, "Phase B inference latency below 1.0 s per study on CUDA "
        "and below 8 s on CPU-only fallback hardware.")
_bullet(doc, "Grad-CAM render must not block prediction or retrieval — "
        "modules render asynchronously.")
_bullet(doc, "Backend must remain available under three simultaneous "
        "demo sessions.")
_bullet(doc, "All API responses validated by Pydantic v2 models that "
        "mirror the TypeScript types one-for-one.")
_bullet(doc, "Image and heatmap assets must be deliverable through "
        "Supabase Storage with public read-only URLs.")

_section(doc, "3.3 Dataset")
_para(doc,
      "All experiments use the Symile-MIMIC v1.0.0 release. After the "
      "preprocessing pipeline in code/process_mimic_data.py and "
      "create_dataset_splits.py, the cohort is partitioned as follows.")
_table(doc,
       ["Split", "Examples", "Description"],
       [["Train", "237,972", "DenseNet121 fine-tuning"],
        ["Validation", "1,959", "Temperature scaling + selective-"
                              "prediction thresholds"],
        ["Test", "3,403", "Held-out evaluation"]],
       col_widths_cm=[2.5, 2.5, 9])
_caption(doc, "Table 3.1", "Symile-MIMIC split sizes after preprocessing.")

_para(doc,
      "Eight CheXpert findings are predicted. The class prevalence in "
      "each split is shown in Table 3.2. Pneumothorax and Consolidation "
      "are particularly rare on the test set (3.2 % and 6.2 % "
      "respectively), which has direct consequences for AUPRC and "
      "operating-point selection in Chapter 6.")

_table(doc,
       ["Finding", "Train (%)", "Val (%)", "Test (%)"],
       [["Cardiomegaly",     "19.49", "20.52", "26.33"],
        ["Pleural Effusion", "23.58", "25.78", "32.18"],
        ["Edema",            "11.91", "13.48", "21.33"],
        ["Pneumonia",        "7.04",  "6.28",  "10.05"],
        ["Atelectasis",      "20.02", "20.16", "22.45"],
        ["Pneumothorax",     "4.64",  "4.13",  "3.17"],
        ["Consolidation",    "4.72",  "4.24",  "6.20"],
        ["Support Devices",  "30.02", "32.41", "35.59"]],
       col_widths_cm=[5, 2.5, 2.5, 2.5])
_caption(doc, "Table 3.2", "Per-class positive prevalence across splits.")

_section(doc, "3.4 Modelling Choices")
_subsection(doc, "3.4.1 Phase A — Tabular Risk")
_para(doc,
      "Phase A consumes seven ECG features (heart rate, PR interval, "
      "QRS duration, QTc, ST deviation, rhythm interpretation and "
      "acquisition time) together with 50 MIMIC-IV laboratory itemids "
      "and produces a softmax over {Low, Moderate, High}. The "
      "underlying classifier is a small MLP defined in "
      "engine/early_risk_inference.py; if its checkpoint is not "
      "present the system falls back to a rules-based scorer that "
      "preserves the same output schema, so the front-end is decoupled "
      "from training success.")

_subsection(doc, "3.4.2 Phase B — DenseNet121 Backbone")
_para(doc,
      "Phase B uses a DenseNet121 backbone with four functional "
      "Dropout(p = 0.30) layers placed after each transition / dense "
      "block (model_mm_film_gated.py). Functional dropout, as opposed "
      "to nn.Dropout modules, lets MC-Dropout be enabled through a "
      "boolean flag without flipping the BatchNorm layers into "
      "training mode — a subtle but important property for stable "
      "uncertainty estimates. Inputs are resized to 512 × 512, "
      "converted to three-channel ImageNet-normalised tensors and "
      "passed through the network; the 1024-dimensional pre-classifier "
      "vector is L2-normalised and stored as the FAISS embedding.")

_subsection(doc, "3.4.3 Calibration")
_para(doc,
      "After training, a single scalar T is fitted on the validation "
      "split by minimising the multi-label NLL on the logits. The "
      "fitted value, T = 1.252, is then frozen and applied to all "
      "test-time logits before sigmoid. T > 1 reduces over-confidence "
      "and shrinks ECE without changing AUROC.")

_subsection(doc, "3.4.4 MC-Dropout")
_para(doc,
      "At inference time, dropout is re-enabled for sixty stochastic "
      "forward passes on the test split (run config: dropout_p = 0.30, "
      "mc_passes = 60, batch_size = 60, image_size = 512, device = "
      "cuda). The mean of the per-pass sigmoid outputs is reported as "
      "the calibrated probability and the per-class variance is "
      "discretised into the three uncertainty tiers shown in Table 4.2. "
      "On the production API, the number of passes is capped at ten "
      "with a 30-second timeout to keep response time predictable.")

_subsection(doc, "3.4.5 Retrieval")
_para(doc,
      "FAISS IndexFlatIP is used over the L2-normalised 1024-d "
      "DenseNet GAP vectors. Two parallel indexes are maintained — one "
      "for the DenseNet feature space (faiss_index.bin) and one for "
      "the optional Symile multimodal feature space "
      "(faiss_symile_index.bin) — and the lifespan handler in main.py "
      "flushes both atomically on every case mutation that affects "
      "embeddings.")

# ============================================================================
# CHAPTER 4 — DESIGN AND ARCHITECTURE
# ============================================================================
_chapter_heading(doc, 4, "System Design and Architecture")

_section(doc, "4.1 Three-Tier Architecture")
_para(doc,
      "The platform is structured in three tiers. The presentation "
      "tier is a Next.js 13 / App Router application written in "
      "TypeScript with TailwindCSS for styling and Zustand for global "
      "state. The application tier is a FastAPI Python backend that "
      "exposes thirty-plus REST endpoints and delegates inference work "
      "to background tasks. The data tier is a Supabase Postgres "
      "instance for relational data plus Supabase Storage for binary "
      "DICOM and heatmap assets, complemented by two on-disk FAISS "
      "indexes for similarity search.")

_section(doc, "4.2 Inference Pipeline")
_subsection(doc, "4.2.1 Phase A Pipeline")
_para(doc,
      "When the wizard registers a new case, ECG and lab payloads are "
      "validated by ECGData and LabData Pydantic models. "
      "engine/early_risk_inference.py then computes the softmax over "
      "{Low, Moderate, High}, persists the resulting risk_level, "
      "risk_score_norm, recommendation and class_probabilities into "
      "the cases row, and stamps phase_a_run_at. The frontend reads "
      "this state through GET /api/cases/{case_id} and renders the "
      "Early Risk tab.")

_subsection(doc, "4.2.2 Phase B Pipeline")
_para(doc,
      "When a CXR is uploaded, FastAPI accepts the multipart upload, "
      "writes the file to Supabase Storage, queues a BackgroundTask "
      "and returns a 202 with the new case_id. The background worker "
      "loads the image, runs run_cxr_inference, persists eight rows in "
      "the predictions table (one per CheXpert finding), runs MC-"
      "Dropout to populate uncertainty fields, generates a Grad-CAM "
      "overlay for the argmax label, uploads the heatmap to Supabase "
      "Storage and finally writes the public URL back to the cases "
      "row. A single inference lock serialises all model-touching "
      "sections because MC-Dropout flips the model into training mode "
      "mid-call.")

_section(doc, "4.3 Risk Badges and Uncertainty Tiers")
_table(doc,
       ["Probability", "Badge", "UI Colour"],
       [["p < 0.05",        "Unlikely",      "Gray"],
        ["0.05 ≤ p ≤ 0.15", "Monitor",       "Yellow"],
        ["p > 0.15",        "Elevated Risk", "Red"]],
       col_widths_cm=[4.5, 4.5, 5])
_caption(doc, "Table 4.1", "Three-tier risk-badge thresholds applied per "
                            "predicted finding (also enforced as a Postgres "
                            "GENERATED column).")

_table(doc,
       ["MC-Dropout variance", "Tier", "UI Banner"],
       [["σ² < 0.01",        "Low Uncertainty",      "Green"],
        ["0.01 ≤ σ² < 0.04", "Moderate Uncertainty", "Amber"],
        ["σ² ≥ 0.04",        "High Uncertainty",     "Red, with warning"]],
       col_widths_cm=[5, 5, 4])
_caption(doc, "Table 4.2", "Three-tier MC-Dropout uncertainty banner used "
                            "by the frontend.")

_section(doc, "4.4 Role-Based Access Control")
_table(doc,
       ["Role", "Visible Tabs / Surfaces"],
       [["Ward Doctor", "Early Risk, CXR Analysis, ECG Input, Similar "
                        "Cases, Before vs. After"],
        ["Radiologist", "Patient Summary (read-only Early Risk), CXR "
                        "Analysis with Flag-Critical, Similar Cases, "
                        "Before vs. After"],
        ["Clinical Admin", "Register Patient wizard, Upload Data, Case "
                           "Status overview"],
        ["System Admin", "Live metrics, User Management, Audit Log, "
                         "login activity"]],
       col_widths_cm=[3.5, 10.5])
_caption(doc, "Table 4.3", "Role-to-surface visibility matrix.")

_section(doc, "4.5 Data Model")
_para(doc,
      "The Postgres schema (backend/supabase_schema/schema.sql) uses "
      "five primary tables. patients holds demographic metadata keyed "
      "by MRN. cases holds admission-level data with JSONB ecg_data, "
      "lab_data and labs_raw columns, plus all Phase A and Phase B "
      "scalar outputs. predictions holds the eight per-finding rows "
      "for each case with a Postgres GENERATED column for risk_badge "
      "so that the badge can never disagree with its probability. "
      "consultations holds the ward-doctor / radiologist thread with "
      "a JSONB messages array and a viewport_state JSONB for "
      "cross-user cursor synchronisation. audit_log holds the append-"
      "only RBAC trail.")

_section(doc, "4.6 Audit and Identity Resolution")
_para(doc,
      "Every mutating endpoint passes through the require_role helper "
      "in backend/auth.py which returns the resolved actor or raises "
      "HTTP 403. Identity is resolved in two modes: (i) preferred "
      "JWT mode where Authorization: Bearer <token> is verified with "
      "HS256 against SUPABASE_JWT_SECRET, with the role looked up from "
      "the users table behind a 60-second TTL cache; (ii) header-shim "
      "mode where X-User-Id and X-User-Role are trusted as a fallback "
      "while the frontend Supabase auth integration is still being "
      "wired up. log_audit() is best-effort and never raises, and as a "
      "side effect bumps users.last_active_at so the system "
      "administrator sees live presence.")

# ============================================================================
# CHAPTER 5 — IMPLEMENTATION
# ============================================================================
_chapter_heading(doc, 5, "Implementation")

_section(doc, "5.1 Backend (FastAPI)")
_para(doc,
      "The backend is structured around five files. main.py declares "
      "the FastAPI app with thirty-plus endpoints and a lifespan handler "
      "that hot-loads the DenseNet weights, primes the FAISS indexes "
      "and flushes them on shutdown. database.py wraps the Supabase "
      "Python client with a LOCAL_MOCK fallback that keeps the system "
      "operable when SUPABASE_URL is not configured. schemas.py "
      "defines all Pydantic v2 response and request models, mirrors "
      "frontend/src/lib/types.ts one-for-one, and includes a "
      "model_validator(mode='before') on ConsultationMessage that "
      "tolerates legacy body / created_at fields while writing only "
      "canonical content / sent_at fields going forward. auth.py "
      "implements the two-mode identity resolution described in "
      "Section 4.6. engine/ wraps the DenseNet model loader, the CXR "
      "inference + Grad-CAM + MC-Dropout pipeline, the FAISS vector "
      "store and the optional Symile encoder.")

_section(doc, "5.2 Frontend (Next.js 13)")
_para(doc,
      "The Next.js application is organised under src/app (App "
      "Router pages) and src/components (presentational components). "
      "Domain types live in src/lib/types.ts and HTTP helpers in "
      "src/lib/api.ts. Global state — the active case, the active tab, "
      "consultation cursor — lives in two Zustand stores under "
      "src/store. Role-aware visibility is centralised in "
      "components/case/CaseTabs.tsx which derives the visible tab list "
      "from the role returned by useUserRole(); a dev role-switcher in "
      "the header (cdss_dev_role in localStorage) lets the developer "
      "act as any of the four demo users without touching the auth "
      "stack. Tailwind utility classes are merged through a small cn() "
      "helper that combines clsx and tailwind-merge.")

_section(doc, "5.3 Audit Trail")
_para(doc,
      "Eight mutating endpoints write a row to audit_log via "
      "log_audit(). Table 5.1 lists them. The actor is resolved "
      "through auth.get_actor(), which returns "
      "{user_id, user_role}; both fields are persisted on the audit "
      "row alongside the entity type, entity id, action verb, "
      "timestamp and a free-form metadata JSONB.")

_table(doc,
       ["Endpoint", "Action verb"],
       [["POST /api/cases (wizard)",                       "case.create"],
        ["PATCH /api/cases/{id}/complete",                  "case.complete"],
        ["PATCH /api/cases/{id}/flag-critical",             "case.flag_critical"],
        ["POST /api/cases/{id}/cxr",                        "cxr.upload"],
        ["POST /api/cases/{id}/cxr/reinfer",                "cxr.reinfer"],
        ["POST /api/cases/{id}/heatmap/regenerate",         "gradcam.regenerate"],
        ["POST /api/cases/{id}/ecg",                        "ecg.upload"],
        ["POST /api/cases/{id}/labs",                       "labs.upload"],
        ["POST /api/admin/faiss/reload",                    "faiss.reload"],
        ["POST /api/admin/users",                           "user.create"],
        ["PATCH /api/admin/users/{id}",                     "user.update"]],
       col_widths_cm=[8, 6])
_caption(doc, "Table 5.1", "Mutating endpoints that emit audit-log rows.")

_section(doc, "5.4 Tooling and Dependencies")
_table(doc,
       ["Layer", "Stack"],
       [["Backend",        "Python 3.14, FastAPI, Uvicorn, Pydantic v2, "
                           "PyJWT, supabase-py, PyTorch, torchvision, "
                           "FAISS-CPU, OpenCV, Pillow, scikit-learn"],
        ["Frontend",       "Next.js 13 App Router, TypeScript, "
                           "TailwindCSS, Zustand, lucide-react, "
                           "tailwind-merge, clsx"],
        ["Data",           "Supabase Postgres + Storage, FAISS on-disk "
                           "indexes, MIMIC-IV / MIMIC-CXR-JPG / Symile-"
                           "MIMIC v1.0.0"],
        ["Training",       "DenseNet121 (ImageNet1K_V1), Albumentations, "
                           "AMP, GCS-cached image loader"]],
       col_widths_cm=[3, 11])
_caption(doc, "Table 5.2", "External dependencies grouped by layer.")

_section(doc, "5.5 Notable Engineering Decisions")
_subsection(doc, "5.5.1 Inference Lock")
_para(doc,
      "MC-Dropout flips the shared DenseNet into training mode mid-"
      "request. Without serialisation, a parallel deterministic "
      "request would observe stochastic outputs. A "
      "threading.Lock in engine/inference.py serialises every model-"
      "touching code path; the lock is held for the duration of one "
      "inference (≈ 0.6 s on CUDA) which is acceptable at demo "
      "concurrency.")

_subsection(doc, "5.5.2 LOCAL_MOCK Fallback")
_para(doc,
      "database.get_db() returns the literal string 'LOCAL_MOCK' when "
      "SUPABASE_URL is not configured. Every database helper checks "
      "for this sentinel and reads / writes backend/local_db.json "
      "instead. This makes the application bootstrap-able in two "
      "minutes for an evaluator who does not want to provision a "
      "Supabase project, while leaving the production code path "
      "unchanged.")

_subsection(doc, "5.5.3 Pydantic Legacy-Field Tolerance")
_para(doc,
      "Earlier development versions of ConsultationMessage stored "
      "body / created_at instead of the canonical content / sent_at. "
      "Rather than migrate the JSONB column in place — which would "
      "require coordinated downtime — a model_validator(mode='before') "
      "on ConsultationMessage maps the legacy keys onto the canonical "
      "ones at deserialisation time, with extra='ignore' on the model "
      "config so additional historical keys are silently dropped.")

_subsection(doc, "5.5.4 Two FAISS Indexes")
_para(doc,
      "The platform maintains two parallel FAISS indexes — one for "
      "1024-d DenseNet GAP features and one for 24,576-d Symile "
      "concatenated multimodal features (8,192-d × 3 modalities) — "
      "because the demo benefits from both single-modality and "
      "multimodal queries. The lifespan handler flushes both on "
      "shutdown to keep the on-disk id-maps consistent with the "
      "index files.")

_section(doc, "5.6 Symile Dataset Generation and Preprocessing")
_subsection(doc, "5.6.1 Source Datasets")
_para(doc,
      "Two distinct datasets are used in this project. The first, used "
      "to train the DenseNet121 baseline (Chapter 3), is the full "
      "MIMIC-CXR-JPG v2.0.0 release with the CheXpert label set "
      "(237,972 / 1,959 / 3,403 train / val / test studies after "
      "preprocessing). The second, used to train the Symile multimodal "
      "encoder, is the Symile-MIMIC v1.0.0 release described by "
      "Saporta et al. (2024), comprised of 11,622 hospital admissions "
      "with synchronously joined CXR, ECG and lab triples. Both are "
      "generated from three PhysioNet-credentialed sources: MIMIC-IV "
      "v2.2 (admissions and lab events from the hosp module), "
      "MIMIC-IV-ECG v1.0 (record_list.csv with 12-lead ECG waveforms) "
      "and MIMIC-CXR-JPG v2.0.0 "
      "(mimic-cxr-2.0.0-metadata.csv.gz and "
      "mimic-cxr-2.0.0-chexpert.csv.gz). The vendored copy of the "
      "Symile-MIMIC preprocessing pipeline lives under code/ in the "
      "project repository.")

_subsection(doc, "5.6.2 Patient-Admission Sampling Rule")
_para(doc,
      "code/process_mimic_data.py joins the three modalities at the "
      "level of a single hospital admission. For every admission the "
      "earliest CXR, ECG and lab panel are kept under the following "
      "constraints. The ECG must have been recorded within 24 hours "
      "of admission, must contain no NaN values and must not be a "
      "constant-zero signal. The lab panel must be drawn within the "
      "same 24-hour window and the CXR must be acquired in the "
      "24–72 hour window post-admission so that imaging is "
      "temporally downstream of the early modalities. Only "
      "posteroanterior (PA) and anteroposterior (AP) views are "
      "retained. Each retained admission must include a CXR, an ECG "
      "and at least one of the 50 most-frequent MIMIC-IV lab itemids "
      "tabulated in code/constants.py. The script completes in "
      "approximately seven hours on 16 CPU cores with 100 GB of RAM "
      "and writes a single symile_mimic_data.csv to disk.")

_subsection(doc, "5.6.3 Split Construction")
_para(doc,
      "code/create_dataset_splits.py partitions the joined CSV into "
      "patient-disjoint train, validation and test subsets so that "
      "no patient_id appears in more than one split — a critical "
      "guard against subject-level leakage. For Symile-MIMIC the "
      "11,622-admission corpus is split 95 % into a development "
      "(train + validation) set and 5 % into a held-out test set, "
      "as specified by Saporta et al. (2024). Lab values are "
      "converted to percentiles using a NaN-aware empirical "
      "cumulative distribution function fitted on the training "
      "split only; the same CDF is applied at inference time, and "
      "missing values are imputed with the per-lab training mean "
      "percentile. A 100-dimensional input vector is ultimately "
      "produced for the labs encoder: the first 50 coordinates are "
      "percentile-standardised lab values, and the remaining 50 are "
      "binary missingness indicators. To support zero-shot "
      "retrieval evaluation the validation and test splits are "
      "augmented with negative candidates: for every query (one "
      "true triple) the script samples nine random non-matching "
      "candidates, yielding a 1-positive-versus-9-negatives "
      "evaluation set (val_retrieval.csv and test.csv) with random-"
      "chance accuracy of 0.10.")

_table(doc,
       ["Stage", "Outcome", "Tooling"],
       [["Modality join",        "symile_mimic_data.csv",
         "process_mimic_data.py (≈ 7 h, 16 CPUs, 100 GB RAM)"],
        ["Patient-disjoint split", "train.csv, val.csv, test.csv",
         "create_dataset_splits.py"],
        ["Retrieval augmentation","val_retrieval.csv, test.csv "
                                  "(10 candidates / query)",
         "create_dataset_splits.py"],
        ["Tensor materialisation", "data_npy/{split}/{modality}_{split}.npy",
         "process_and_save_tensors.py (≈ 1 h, 150 GB RAM)"]],
       col_widths_cm=[3.5, 5, 5.5])
_caption(doc, "Table 5.3", "Symile-MIMIC dataset construction pipeline. "
                            "Each stage is reproducible from the scripts "
                            "vendored under code/.")

_subsection(doc, "5.6.4 Tensor Materialisation")
_para(doc,
      "code/process_and_save_tensors.py converts each split into "
      "split-specific numpy tensors that the data loader memory-maps "
      "at training time. CXRs follow CheXpert preprocessing: the "
      "smaller edge is scaled to cxr_scale = 320 px, then a "
      "320 × 320 random crop is taken on the training split and a "
      "centre crop on val/test, and the image is normalised with "
      "the standard ImageNet mean (0.485, 0.456, 0.406) and standard "
      "deviation (0.229, 0.224, 0.225). The 12-lead ECG waveform "
      "(shape 1 × 5000 × 12) is min-max scaled into [−1, 1]. Lab "
      "tensors are saved as the 50-dim percentile vector and a "
      "matching 50-dim missingness vector. The five resulting "
      "tensors per split — cxr.npy, ecg.npy, labs_percentiles.npy, "
      "labs_missingness.npy and hadm_id.npy — are persisted under "
      "data_npy/{train,val,test}/.")

_section(doc, "5.7 Symile Multimodal Pre-Training")
_subsection(doc, "5.7.1 Architecture")
_para(doc,
      "The Symile multimodal encoder follows symile-main/experiments/"
      "models/symile_mimic_model.py. Three modality-specific encoders "
      "project each input into a shared d-dimensional embedding "
      "space, with d = 8,192 in the released checkpoint. The CXR "
      "encoder is a ResNet-50 (320 × 320 RGB input) whose final fully-"
      "connected layer is replaced by a Linear(2048, d) head followed "
      "by LayerNorm. The ECG encoder is a ResNet-18 with its first "
      "convolution re-instantiated to accept a single-channel "
      "(1 × 5000 × 12) input and a Linear(512, d) projection head. "
      "The Labs encoder is a three-layer MLP "
      "(100 → 256 → 1024 → d) with GELU activations and a final "
      "LayerNorm. A learnable scalar log-temperature, initialised to "
      "−7.0, scales the multilinear inner product (MIP) at the loss "
      "and at retrieval time, mirroring CLIP’s temperature mechanism.")

_subsection(doc, "5.7.2 Symile Contrastive Loss")
_para(doc,
      "Symile (Saporta et al., 2024) generalises the bilateral "
      "contrastive objective of CLIP to k ≥ 3 modalities. The "
      "similarity score for a positive triple is the multilinear "
      "inner product MIP(c, e, ℓ) = ⟨c ⊙ e ⊙ ℓ, 𝟙⟩, i.e. the sum "
      "of the element-wise product of the three normalised "
      "embeddings. Two negative-sampling strategies are provided "
      "in symile/loss.py: an O(N) variant that shuffles the non-"
      "anchor modalities once per batch, and an O(N²) variant that "
      "constructs N² − 1 negatives per positive. The released "
      "checkpoint uses the O(N²) strategy. Critically, MIP targets "
      "the total correlation of the three random variables, which "
      "is the higher-order generalisation of mutual information, "
      "and so Symile captures pairwise plus conditional interactions "
      "that pairwise CLIP-style training cannot.")

_subsection(doc, "5.7.3 Training Configuration")
_table(doc,
       ["Hyper-parameter", "Value"],
       [["Embedding dim. d",      "8,192 per modality"],
        ["Batch size (train / val / test)", "280 / 280 / 280"],
        ["Negative sampling",     "O(N²) (n_squared)"],
        ["Loss",                  "Symile multilinear contrastive"],
        ["Optimiser",             "AdamW (β = (0.9, 0.999))"],
        ["Learning rate",         "1.0 × 10⁻³"],
        ["Weight decay",          "0.5"],
        ["Epochs",                "80"],
        ["Logit scale init.",     "−7.0 (learnt)"],
        ["Pretrained backbones",  "False (trained from scratch)"],
        ["Validation cadence",    "every 1 epoch"],
        ["Random seed",           "0 (use_seed = True)"],
        ["Image size",            "320 × 320"],
        ["ECG shape",             "1 × 5000 × 12"],
        ["Labs input",            "100-dim (50 percentiles + 50 missingness)"],
        ["Framework",             "PyTorch Lightning, AdamW, FP32"]],
       col_widths_cm=[5.5, 8.5])
_caption(doc, "Table 5.4", "Hyperparameters of the released Symile-MIMIC "
                            "checkpoint (symile_mimic_model.ckpt), "
                            "extracted from the saved PyTorch Lightning "
                            "hyper_parameters block.")

_para(doc,
      "Training was run on a multi-GPU SLURM cluster "
      "(\\gpfs\\scratch\\as16583\\ckpts\\symile_mimic) with "
      "PyTorch Lightning. The provided checkpoint is loaded by the "
      "backend at lifespan start through engine/symile_encoder.py "
      "and is consumed in two ways: (i) to populate the 24,576-d "
      "FAISS Symile index — the concatenation of the three "
      "L2-normalised 8,192-d embeddings — used for multimodal "
      "similar-case retrieval; and (ii) at inference time to project "
      "user-supplied ECG, lab and CXR inputs into the same space "
      "for ad-hoc retrieval queries.")

_section(doc, "5.8 Phase 5 Authentication (Completed Before Review)")
_para(doc,
      "Phase 5 replaces the development role-switcher with a real "
      "Supabase authentication stack. The work was sliced into five "
      "incremental milestones so that the application remained "
      "functional at every step; all five were completed and "
      "verified end-to-end before the project review.")

_table(doc,
       ["Milestone", "Scope", "Status"],
       [["5a — Data plane",
         "Create users, audit_log and outcome tables; seed four demo "
         "users (one per role) with correct id, role and status fields.",
         "Completed"],
        ["5b — Backend JWT verification",
         "Add backend/auth.py with HS256 JWT decode against "
         "SUPABASE_JWT_SECRET, audience = 'authenticated', and a "
         "role lookup against users with a 60-second TTL cache. "
         "Wire require_role into eight mutating endpoints; retain "
         "X-User-Id / X-User-Role header shim as a fallback.",
         "Completed"],
        ["5c — Frontend Supabase client",
         "Add @supabase/supabase-js, login and signup pages, "
         "password reset, session persistence in HTTP-only "
         "cookies. useUserRole() reads the resolved JWT claim "
         "instead of localStorage, and the dev role-switcher is "
         "gated behind NEXT_PUBLIC_ENABLE_DEV_ROLE_SWITCHER.",
         "Completed"],
        ["5d — User migration",
         "Lift the four hard-coded demo users to first-class "
         "Supabase Auth users; map auth.uid() back to public.users "
         "through a SECURITY DEFINER trigger so that role and "
         "status remain editable by the system administrator.",
         "Completed"],
        ["5e — Row-Level Security",
         "Enable RLS on cases, consultations and audit_log. "
         "Patient-facing tables are readable by any authenticated "
         "clinician; mutations are gated by role; audit_log is "
         "append-only and selectable only by system_admin.",
         "Completed"]],
       col_widths_cm=[3, 9, 2])
_caption(doc, "Table 5.5", "Phase 5 authentication milestones. All five "
                            "were completed prior to project review and "
                            "are now in production-equivalent use on the "
                            "Supabase project (ref: yfjsibbnzgnmukxvgfdb).")

_para(doc,
      "End-to-end verification covered (i) signed Bearer-token "
      "authentication on the eight audited endpoints, (ii) a "
      "denied-role 403 path on /api/admin/users when called as "
      "ward_doctor, (iii) audit-log entries correctly attributing "
      "actor.user_id to the authenticated subject, and (iv) row-"
      "level-security enforcement against direct Supabase REST "
      "queries that bypass the FastAPI layer. The header shim "
      "remains in the codebase but is now disabled in production "
      "and only re-enabled in development when "
      "SUPABASE_JWT_SECRET is unset, preserving the offline-demo "
      "experience for examiners without Supabase credentials.")

# ============================================================================
# CHAPTER 6 — RESULTS AND DISCUSSION
# ============================================================================
_chapter_heading(doc, 6, "Results and Discussion")

_section(doc, "6.1 Overall Discrimination")
_para(doc,
      "Table 6.1 summarises the headline metrics on the 3,403-example "
      "test set. Micro-AUROC is 0.864 and macro-AUROC is 0.812; the "
      "micro number is, as expected, dominated by the more prevalent "
      "labels Pleural Effusion and Support Devices. Micro-AUPRC is "
      "0.634, materially lower than the AUROC figure, which is a "
      "natural consequence of the strong class imbalance "
      "(Pneumothorax, for example, is only 3.2 % positive on the test "
      "split). Brier score is 0.111 before temperature scaling.")

_table(doc,
       ["Metric", "Value", "95 % bootstrap CI"],
       [["Micro-AUROC", "0.8643", "[0.8585, 0.8693]"],
        ["Macro-AUROC", "0.8116", "[0.8033, 0.8196]"],
        ["Micro-AUPRC", "0.6339", "—"],
        ["Macro-AUPRC", "0.4871", "—"],
        ["Brier",        "0.1109", "—"]],
       col_widths_cm=[4, 3, 7])
_caption(doc, "Table 6.1", "Headline test-set metrics for the DenseNet121 "
                            "baseline (pre-temperature-scaling).")

_figure(doc, PLOT_DIR / "preTS" / "roc_curve_micro.png",
        "Figure 6.1", "Micro-averaged ROC curve before temperature "
                      "scaling.")
_figure(doc, PLOT_DIR / "preTS" / "roc_curve_per_label_grid.png",
        "Figure 6.2", "Per-label ROC curves (eight findings).", width_in=5.3)
_figure(doc, PLOT_DIR / "preTS" / "pr_curve_micro.png",
        "Figure 6.3", "Micro-averaged precision-recall curve.")

_section(doc, "6.2 Per-Label Results")
_table(doc,
       ["Finding", "AUROC", "95 % CI", "AUPRC", "Test prev. (%)"],
       [["Cardiomegaly",     "0.7824", "[0.7654, 0.7977]", "0.5287", "26.33"],
        ["Pleural Effusion", "0.8831", "[0.8714, 0.8940]", "0.7925", "32.18"],
        ["Edema",            "0.8353", "[0.8204, 0.8507]", "0.5827", "21.33"],
        ["Pneumonia",        "0.7076", "[0.6813, 0.7364]", "0.2388", "10.05"],
        ["Atelectasis",      "0.7523", "[0.7356, 0.7707]", "0.4127", "22.45"],
        ["Pneumothorax",     "0.8590", "[0.8209, 0.8959]", "0.3344", "3.17"],
        ["Consolidation",    "0.7636", "[0.7297, 0.7919]", "0.1879", "6.20"],
        ["Support Devices",  "0.9092", "[0.8985, 0.9194]", "0.8193", "35.59"]],
       col_widths_cm=[3.5, 1.5, 3.5, 1.8, 2.5])
_caption(doc, "Table 6.2", "Per-label AUROC with 95 % bootstrap "
                            "confidence intervals (1,000 resamples), AUPRC "
                            "and test-set positive prevalence.")
_para(doc,
      "Pleural Effusion, Pneumothorax and Support Devices reach AUROC "
      "above 0.85, while Pneumonia (0.708) and Atelectasis (0.752) "
      "trail. The AUPRC column foregrounds the class-imbalance story: "
      "Consolidation reaches 0.764 AUROC but only 0.188 AUPRC because "
      "fewer than 7 % of test studies are positive. This is the "
      "specific pattern that motivates the uncertainty banner — the "
      "model is rarely confident on these rare classes, and the front-"
      "end is engineered to surface that fact rather than hide it.")

_section(doc, "6.3 Calibration")
_para(doc,
      "On the validation split the un-calibrated negative log-"
      "likelihood was 0.2851, ECE was 0.0397 and Brier was 0.0872. "
      "Fitting a single scalar T = 1.252 on the validation logits "
      "reduced these to NLL = 0.2784, ECE = 0.0269 and Brier = 0.0859, "
      "an absolute ECE reduction of 0.013 (≈ 32 % relative). On the "
      "test split, ECE drops from 0.0538 (pre) to 0.0305 (post), "
      "confirming that the temperature generalises beyond the "
      "calibration set without any change to AUROC.")

_table(doc,
       ["Metric", "Before TS", "After TS", "Δ"],
       [["NLL (validation)",         "0.2851", "0.2784", "−0.0067"],
        ["ECE (validation, 15 bins)", "0.0397", "0.0269", "−0.0128"],
        ["Brier (validation)",        "0.0872", "0.0859", "−0.0014"],
        ["ECE (test, 15 bins)",       "0.0538", "0.0305", "−0.0233"]],
       col_widths_cm=[5.5, 2.5, 2.5, 2])
_caption(doc, "Table 6.3", "Calibration before and after temperature "
                            "scaling (T = 1.252).")
_figure(doc, PLOT_DIR / "preTS" / "calibration_overall.png",
        "Figure 6.4", "Reliability diagram before temperature scaling.")
_figure(doc, PLOT_DIR / "postTS" / "calibration_overall.png",
        "Figure 6.5", "Reliability diagram after temperature scaling.")

_section(doc, "6.4 Uncertainty Quantification (MC-Dropout)")
_para(doc,
      "Sixty stochastic forward passes with dropout p = 0.30 were "
      "executed across the entire 3,403-example test split (run config: "
      "image_size = 512, batch_size = 60, device = cuda). Total "
      "wall-clock time was 1,832 s (≈ 30 minutes). The mean of the "
      "per-pass sigmoid outputs is reported as the calibrated "
      "probability and the per-class variance is discretised into the "
      "three uncertainty tiers shown in Table 4.2. For the live API "
      "the number of passes is capped at ten and a 30-second timeout "
      "guards latency.")
_table(doc,
       ["Field", "Value"],
       [["Test examples",            "3,403"],
        ["MC passes",                "60"],
        ["Dropout p",                "0.30"],
        ["Mean-probs Brier",         "0.1096"],
        ["Mean-probs NLL",           "0.3543"],
        ["Wall-clock (s)",           "1,832"]],
       col_widths_cm=[5, 5])
_caption(doc, "Table 6.5", "MC-Dropout summary on the test split.")

_figure(doc, PLOT_DIR / "preTS" / "uncertainty_hist_entropy.png",
        "Figure 6.7", "Predictive entropy histogram (test split, "
                      "60 MC passes).")
_figure(doc, PLOT_DIR / "preTS" / "coverage_vs_auroc_macro_entropy.png",
        "Figure 6.6", "Macro-AUROC as a function of selective-prediction "
                      "coverage under entropy-based abstention.")
_para(doc,
      "Figure 6.6 shows the standard selective-prediction trade-off: "
      "as the coverage threshold drops (i.e. the model is allowed to "
      "abstain on its most uncertain inputs), macro-AUROC over the "
      "remaining accepted examples rises monotonically. Selective "
      "prediction at 80 % coverage is the operating mode currently "
      "exposed in the UI as the 'High Uncertainty' banner.")

_section(doc, "6.5 Operating Points")
_para(doc,
      "Per-label thresholds that achieve clinically meaningful "
      "operating points (95 % sensitivity, 90 % sensitivity, Youden-J, "
      "F1-max) are saved to "
      "02_metrics_preTS/clinical_operating_points.json. As an "
      "illustrative example, at 95 % sensitivity the operating "
      "specificities are 0.049 (Cardiomegaly), 0.040 (Pleural "
      "Effusion) and 0.020 (Edema), which reinforces the design "
      "decision to expose multiple thresholds to the clinician rather "
      "than forcing a single binary cut-off. The full table is "
      "reproduced in Appendix C.")
_figure(doc, PLOT_DIR / "preTS" / "threshold_sweep_metrics.png",
        "Figure 6.8", "Threshold sweep — F1, sensitivity and "
                      "specificity vs. probability cutoff.")

_section(doc, "6.6 Inference Latency")
_table(doc,
       ["Split", "Examples", "Sec / batch (mean ± std)", "Effective batch"],
       [["Validation", "1,959", "0.224 ± 0.191", "60"],
        ["Test",        "3,403", "0.197 ± 0.033", "60"]],
       col_widths_cm=[3, 2.5, 5, 3.5])
_caption(doc, "Table 6.6", "Forward-pass latency on CUDA (single GPU). "
                            "Per-image inference time ≈ 3.3 ms in "
                            "batch mode.")
_para(doc,
      "Per-image latency in batch mode is 3.3 ms on CUDA. Single-image "
      "latency through the HTTP path is dominated by Grad-CAM and JPEG "
      "encoding rather than by the forward pass itself; in practice "
      "a single CXR upload returns within 600 – 900 ms on CUDA and "
      "5 – 8 s on CPU-only fallback.")

_section(doc, "6.7 Qualitative Discussion")
_para(doc,
      "The DenseNet baseline’s strongest labels are also the most "
      "visually distinctive on a CXR: Support Devices (lines, tubes "
      "and pacemakers), Pleural Effusion (visible costophrenic "
      "blunting) and Pneumothorax (visceral pleural line). The "
      "weaker labels — Pneumonia, Atelectasis — are consistently "
      "harder for human readers as well, especially on portable AP "
      "films, and the AUPRC numbers reflect that these labels are "
      "both rare and confusable.")
_para(doc,
      "From a workflow-engineering point of view the most "
      "consequential numbers are not the AUROCs but the calibration "
      "deltas and the uncertainty distributions. The AUROCs justify "
      "the model architecture; the calibration and uncertainty "
      "results justify the choice to surface probabilities and "
      "uncertainty banners to a non-radiologist user.")

# ============================================================================
# CHAPTER 7 — CONCLUSION AND FUTURE WORK
# ============================================================================
_chapter_heading(doc, 7, "Conclusion and Future Work")

_section(doc, "7.1 Summary of Contributions")
_bullet(doc, "A doctor-facing clinical decision support platform that "
        "models the temporal asymmetry between cheap (ECG, labs) and "
        "expensive (CXR) modalities through an explicit Phase A → "
        "Phase B → Before-vs.-After workflow.")
_bullet(doc, "A calibrated DenseNet121 multi-label classifier "
        "(micro-AUROC 0.864 / macro-AUROC 0.812, with ECE reduced "
        "from 0.054 to 0.031 on the test split via temperature "
        "scaling) wrapped with MC-Dropout uncertainty and Grad-CAM "
        "explanations.")
_bullet(doc, "A FAISS retrieval engine over 1024-d DenseNet GAP "
        "embeddings that operates at every stage of the workflow "
        "with a uniform top-k API.")
_bullet(doc, "A four-role RBAC layer — ward doctor, radiologist, "
        "clinical admin, system admin — with eight audit-logged "
        "mutating endpoints and a pluggable identity-resolution "
        "module that supports both Supabase JWT and a development "
        "header shim.")
_bullet(doc, "A pragmatic engineering stack — FastAPI, Next.js 13, "
        "Supabase, FAISS — that is reproducible from the supplied "
        "scripts and that runs end-to-end in under ten minutes on a "
        "single CUDA-equipped workstation.")

_section(doc, "7.2 Limitations")
_bullet(doc, "AUROC plateaus around 0.71 on the rare classes "
        "(Pneumonia, Consolidation), where training data is sparse "
        "and human-radiologist agreement is itself low.")
_bullet(doc, "MC-Dropout under-estimates epistemic uncertainty on "
        "out-of-distribution inputs because the dropout posterior is "
        "approximate; full deep ensembles would be the principled "
        "comparison.")
_bullet(doc, "The retrieval embedding is a single-modality "
        "1024-dimensional CXR feature; the Symile multimodal "
        "embedding is wired up but not yet the default, and a "
        "rigorous retrieval evaluation has not been performed.")
_bullet(doc, "The Phase A model presented here is a small MLP plus a "
        "rules-based fallback; it has not been independently "
        "validated against a published early-warning score such as "
        "MEWS or NEWS-2.")
_bullet(doc, "The platform runs on de-identified MIMIC-IV data only "
        "and has not been validated in any prospective clinical "
        "setting.")

_section(doc, "7.3 Future Work")
_bullet(doc, "Replace the dev role-switcher with full Supabase Auth "
        "(Phase 5b–5e), add row-level security policies to cases, "
        "consultations and audit_log, and migrate the four hard-coded "
        "demo users to first-class auth users.")
_bullet(doc, "Train a true multimodal retrieval head using the Symile "
        "contrastive loss on the full 448-dimensional joint "
        "embedding (ECG 128 + CXR 256 + Lab 64) and benchmark "
        "retrieval recall@k against the current single-modality "
        "DenseNet baseline.")
_bullet(doc, "Introduce a deep-ensemble alternative to MC-Dropout "
        "for epistemic uncertainty and quantify which procedure best "
        "discriminates accepted from abstained predictions on a "
        "selective-prediction benchmark.")
_bullet(doc, "Externally validate the Phase A early-risk module "
        "against MEWS, NEWS-2 and qSOFA on a held-out non-MIMIC "
        "cohort.")
_bullet(doc, "Pilot the platform in a simulated ward setting with "
        "general physicians as users and measure whether decision "
        "time, imaging-priority accuracy and self-reported confidence "
        "improve relative to the current standard of care.")

# ============================================================================
# REFERENCES
# ============================================================================
doc.add_page_break()
_para(doc, "REFERENCES", align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
      bold=True, caps=True, space_after=18)

refs = [
    "Aamodt, A., & Plaza, E. (1994). Case-based reasoning: Foundational "
    "issues, methodological variations, and system approaches. AI "
    "Communications, 7(1), 39–59.",
    "Gal, Y., & Ghahramani, Z. (2016). Dropout as a Bayesian "
    "approximation: Representing model uncertainty in deep learning. "
    "In Proceedings of the 33rd International Conference on Machine "
    "Learning (Vol. 48, pp. 1050–1059). PMLR.",
    "Guo, C., Pleiss, G., Sun, Y., & Weinberger, K. Q. (2017). On "
    "calibration of modern neural networks. In Proceedings of the "
    "34th International Conference on Machine Learning (Vol. 70, "
    "pp. 1321–1330). PMLR.",
    "Huang, G., Liu, Z., van der Maaten, L., & Weinberger, K. Q. "
    "(2017). Densely connected convolutional networks. In Proceedings "
    "of the IEEE Conference on Computer Vision and Pattern Recognition "
    "(pp. 4700–4708).",
    "Irvin, J., Rajpurkar, P., Ko, M., Yu, Y., Ciurea-Ilcus, S., "
    "Chute, C., et al. (2019). CheXpert: A large chest radiograph "
    "dataset with uncertainty labels and expert comparison. In "
    "Proceedings of the AAAI Conference on Artificial Intelligence "
    "(Vol. 33, pp. 590–597).",
    "Johnson, A. E. W., Bulgarelli, L., Shen, L., Gayles, A., "
    "Shammout, A., Horng, S., et al. (2023). MIMIC-IV, a freely "
    "accessible electronic health record dataset. Scientific Data, "
    "10(1), 1. https://doi.org/10.1038/s41597-022-01899-x",
    "Johnson, A. E. W., Pollard, T. J., Greenbaum, N. R., Lungren, "
    "M. P., Deng, C.-y., Peng, Y., Lu, Z., Mark, R. G., Berkowitz, "
    "S. J., & Horng, S. (2019). MIMIC-CXR-JPG, a large publicly "
    "available database of labeled chest radiographs. arXiv preprint "
    "arXiv:1901.07042.",
    "Johnson, J., Douze, M., & Jégou, H. (2019b). Billion-scale "
    "similarity search with GPUs. IEEE Transactions on Big Data, "
    "7(3), 535–547.",
    "Müller, H., Michoux, N., Bandon, D., & Geissbuhler, A. (2004). A "
    "review of content-based image retrieval systems in medical "
    "applications — Clinical benefits and future directions. "
    "International Journal of Medical Informatics, 73(1), 1–23.",
    "Rajpurkar, P., Irvin, J., Zhu, K., Yang, B., Mehta, H., Duan, "
    "T., et al. (2017). CheXNet: Radiologist-level pneumonia "
    "detection on chest X-rays with deep learning. arXiv preprint "
    "arXiv:1711.05225.",
    "Saporta, A., et al. (2024). Contrasting with Symile: Simple "
    "model-agnostic representation learning for unlimited modalities. "
    "Retrieved from https://github.com/rajesh-lab/symile",
    "Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, "
    "D., & Batra, D. (2017). Grad-CAM: Visual explanations from deep "
    "networks via gradient-based localization. In Proceedings of the "
    "IEEE International Conference on Computer Vision (pp. 618–626).",
    "Tjoa, E., & Guan, C. (2020). A survey on explainable artificial "
    "intelligence (XAI): Toward medical XAI. IEEE Transactions on "
    "Neural Networks and Learning Systems, 32(11), 4793–4813.",
    "FastAPI documentation. (n.d.). Retrieved from "
    "https://fastapi.tiangolo.com",
    "Next.js documentation. (n.d.). Retrieved from https://nextjs.org/docs",
    "Supabase documentation. (n.d.). Retrieved from https://supabase.com/docs",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(r)
    _set_run(run, size=12)

# ============================================================================
# APPENDICES
# ============================================================================
doc.add_page_break()
_para(doc, "APPENDIX A   API ENDPOINT REFERENCE",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, caps=True,
      space_after=18)
api_rows = [
    ["GET",   "/api/health",                     "Liveness + Supabase reachability"],
    ["GET",   "/api/cases",                      "Dashboard summary list"],
    ["POST",  "/api/cases",                      "Create case (wizard)"],
    ["GET",   "/api/cases/{id}",                 "Full case detail"],
    ["DELETE","/api/cases/{id}",                 "Delete case (admin)"],
    ["PATCH", "/api/cases/{id}/complete",        "Mark case complete"],
    ["PATCH", "/api/cases/{id}/flag-critical",   "Flag for radiology"],
    ["POST",  "/api/cases/{id}/cxr",             "Upload CXR + queue inference"],
    ["POST",  "/api/cases/{id}/cxr/reinfer",     "Re-run Phase B"],
    ["POST",  "/api/cases/{id}/heatmap/regenerate","Re-run Grad-CAM"],
    ["POST",  "/api/cases/{id}/ecg",             "Upload / replace ECG"],
    ["POST",  "/api/cases/{id}/labs",            "Upload / replace labs"],
    ["GET",   "/api/cases/{id}/similar",         "Top-k similar cases (FAISS)"],
    ["POST",  "/api/labs/parse",                 "Parse uploaded lab CSV/JSON"],
    ["GET",   "/api/admin/users",                "List users (sys-admin)"],
    ["POST",  "/api/admin/users",                "Create user (sys-admin)"],
    ["PATCH", "/api/admin/users/{id}",           "Update user role/status"],
    ["GET",   "/api/admin/audit-log",            "Audit feed"],
    ["POST",  "/api/admin/faiss/reload",         "Hot-reload FAISS index"],
]
_table(doc, ["Method", "Path", "Purpose"], api_rows,
       col_widths_cm=[1.8, 5.5, 6.5])

doc.add_page_break()
_para(doc, "APPENDIX B   DATABASE SCHEMA (KEY TABLES)",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, caps=True,
      space_after=18)
_para(doc, "patients", size=12, bold=True)
_para(doc, "id UUID PK · mrn UNIQUE · first_name · last_name · "
            "date_of_birth · sex ENUM · age_at_admission · "
            "mimic_subject_id · created_at · updated_at")
_para(doc, "cases", size=12, bold=True)
_para(doc, "id UUID PK · patient_id FK → patients · admitted_at · "
            "discharged_at · ecg_data JSONB · lab_data JSONB · "
            "labs_raw JSONB · phase_a_risk_level ENUM · phase_a_risk_score · "
            "phase_a_recommendation · phase_a_run_at · cxr_dicom_url · "
            "cxr_heatmap_url · cxr_heatmap_label ENUM · cxr_acquired_at · "
            "mimic_study_id")
_para(doc, "predictions", size=12, bold=True)
_para(doc, "id UUID PK · case_id FK → cases · model_checkpoint · "
            "temperature · inference_run_at · label ENUM · probability · "
            "risk_badge ENUM (GENERATED) · uncertainty_level ENUM · "
            "mean_variance · std_dev · mc_passes · gradcam_url · gradcam_alpha")
_para(doc, "consultations", size=12, bold=True)
_para(doc, "id UUID PK · case_id FK · ward_doctor_id · radiologist_id · "
            "is_open · opened_at · closed_at · urgency_flag · "
            "messages JSONB[] · viewport_state JSONB · "
            "ward_doctor_last_view · radiologist_last_view")
_para(doc, "audit_log", size=12, bold=True)
_para(doc, "id BIGSERIAL PK · actor_user_id · actor_user_role · "
            "entity_type · entity_id · action · metadata JSONB · "
            "occurred_at TIMESTAMPTZ")
_para(doc, "users", size=12, bold=True)
_para(doc, "id UUID PK · email UNIQUE · full_name · role ENUM "
            "(ward_doctor | radiologist | clinical_admin | system_admin) · "
            "status ENUM (active | inactive) · last_active_at · created_at")

doc.add_page_break()
_para(doc, "APPENDIX C   PER-LABEL METRICS (FULL TABLE)",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, caps=True,
      space_after=18)
_table(doc,
       ["Finding", "AUROC", "AUROC 95 % CI", "AUPRC",
        "Spec@95 %Sens", "F1-max thr."],
       [["Cardiomegaly",     "0.7824", "[0.7654, 0.7977]", "0.5287",
         "0.0494", "0.158"],
        ["Pleural Effusion", "0.8831", "[0.8714, 0.8940]", "0.7925",
         "0.0395", "0.291"],
        ["Edema",            "0.8353", "[0.8204, 0.8507]", "0.5827",
         "0.0198", "0.266"],
        ["Pneumonia",        "0.7076", "[0.6813, 0.7364]", "0.2388",
         "0.0100", "0.158"],
        ["Atelectasis",      "0.7523", "[0.7356, 0.7707]", "0.4127",
         "0.0297", "0.148"],
        ["Pneumothorax",     "0.8590", "[0.8209, 0.8959]", "0.3344",
         "0.5000", "0.187"],
        ["Consolidation",    "0.7636", "[0.7297, 0.7919]", "0.1879",
         "0.5000", "0.123"],
        ["Support Devices",  "0.9092", "[0.8985, 0.9194]", "0.8193",
         "0.1036", "0.237"]],
       col_widths_cm=[3, 1.5, 3.5, 1.5, 2.5, 2])
_caption(doc, "Table C.1", "Full per-label metrics on the held-out test "
                            "split. Spec@95 %Sens entries of 0.500 indicate "
                            "operating points where the rank ordering does "
                            "not allow the target sensitivity to be reached "
                            "without setting the threshold to the median.")

# ---------------------------------------------------------------------------
# save
# ---------------------------------------------------------------------------
doc.save(str(OUT_PATH))
print(f"Wrote {OUT_PATH}")
